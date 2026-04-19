"""Convert docs/case-study.md to docs/case-study.docx for portfolio download."""
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("docs/case-study.md")
DST = Path("docs/case-study.docx")

doc = Document()

# page margins
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

# default font
style = doc.styles["Normal"]
style.font.name = "Helvetica"
style.font.size = Pt(11)

def add_inline(paragraph, text):
    """Parse **bold**, *italic*, `code` and add runs."""
    # tokenize while preserving order
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x8B, 0x1A, 0x1A)
        else:
            paragraph.add_run(part)

lines = SRC.read_text(encoding="utf-8").splitlines()

i = 0
in_code = False
code_buf: list[str] = []
table_buf: list[list[str]] = []
in_table = False

def flush_table():
    global table_buf, in_table
    if not table_buf:
        return
    rows = len(table_buf)
    cols = len(table_buf[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Light Grid Accent 1"
    for r, row_cells in enumerate(table_buf):
        for c, cell_text in enumerate(row_cells):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, cell_text)
            if r == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()
    table_buf = []
    in_table = False

while i < len(lines):
    line = lines[i]

    # code fence
    if line.startswith("```"):
        if in_code:
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_buf))
            run.font.name = "Menlo"
            run.font.size = Pt(9)
            code_buf = []
            in_code = False
        else:
            flush_table()
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # table row detection
    if line.startswith("|") and line.rstrip().endswith("|"):
        # skip separator row
        if re.match(r"^\|\s*[:\-\| ]+\s*\|$", line):
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        table_buf.append(cells)
        in_table = True
        i += 1
        continue
    else:
        if in_table:
            flush_table()

    # horizontal rule
    if line.strip() == "---":
        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(6)
        p_format.space_after = Pt(6)
        run = p.add_run("─" * 40)
        run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        i += 1
        continue

    # headings
    if line.startswith("# "):
        p = doc.add_heading(line[2:].strip(), level=0)
        i += 1
        continue
    if line.startswith("## "):
        p = doc.add_heading(line[3:].strip(), level=1)
        i += 1
        continue
    if line.startswith("### "):
        p = doc.add_heading(line[4:].strip(), level=2)
        i += 1
        continue

    # image
    m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)", line)
    if m:
        alt, src = m.group(1), m.group(2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # try embedding local svg/png; svg may not be supported, so fall back to placeholder line
        if src.startswith("assets/"):
            local = SRC.parent / src
            try:
                doc.add_picture(str(local), width=Cm(14))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                run = p.add_run(f"[ Image: {alt} — {src} ]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        else:
            run = p.add_run(f"[ Image placeholder · {alt} ]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        i += 1
        continue

    # blockquote
    if line.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        run_marker = p.add_run("│ ")
        run_marker.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        add_inline(p, line[2:])
        for run in p.runs[1:]:
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue

    # bullet list
    if re.match(r"^[-*]\s+", line):
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, re.sub(r"^[-*]\s+", "", line))
        i += 1
        continue

    # numbered list
    if re.match(r"^\d+\.\s+", line):
        p = doc.add_paragraph(style="List Number")
        add_inline(p, re.sub(r"^\d+\.\s+", "", line))
        i += 1
        continue

    # italic-only caption line  *foo*
    if line.startswith("*") and line.endswith("*") and not line.startswith("**"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line.strip("*"))
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        i += 1
        continue

    # empty line
    if not line.strip():
        i += 1
        continue

    # plain paragraph
    p = doc.add_paragraph()
    add_inline(p, line)
    i += 1

flush_table()

doc.save(str(DST))
print(f"wrote {DST} ({DST.stat().st_size} bytes)")
